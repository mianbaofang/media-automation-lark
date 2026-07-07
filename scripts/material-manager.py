#!/usr/bin/env python3
"""场景 C：文章链接/图片/PDF/本地文件素材 → AI 分析 → 归档飞书云文档 + 素材表。

用法：
  python3 material-manager.py --url "https://example.com/article"
  python3 material-manager.py --pdf "./report.pdf"
  python3 material-manager.py --image "./cover.png"
  python3 material-manager.py --file "./slides.pptx"   # 自动按扩展名识别
  python3 material-manager.py --queue ./queue.json      # 批量：[{type,value},...]

本地文件（docx/pptx/xlsx/epub/html/音频等）经 markitdown 转成保留结构的
Markdown 后再做 AI 分析；markitdown 缺失时 PDF 回退 PyPDF2、其它格式跳过。
"""
from __future__ import annotations

import argparse
import base64
import hashlib
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import common  # noqa: E402
import file2md  # noqa: E402

MATERIAL_SYSTEM = """你是一个新媒体素材分析师。给定素材内容，输出 JSON：
{
  "title": "素材标题/主题",
  "summary": "100 字以内中文摘要",
  "category": "从[科技,财经,生活,教育,娱乐,其他]选一",
  "tags": ["标签1","标签2"],
  "key_points": ["核心要点1","核心要点2"],
  "action_items": ["可落地的行动项1"],
  "reuse_suggestion": "建议如何使用该素材"
}
只输出 JSON，不要解释；缺失填 null。"""


def fetch_article(url: str):
    if not common.is_safe_url(url):
        common.logger.warning("拒绝不安全 URL（SSRF 防护）: %s", url)
        return "", ""
    html = common.http_get(url)
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    for t in soup(["script", "style", "nav", "footer", "header", "aside"]):
        t.decompose()
    text = "\n".join(line.strip() for line in soup.get_text().splitlines() if line.strip())
    title = soup.title.get_text() if soup.title else url
    return title, text[:6000]


def read_pdf(path: str):
    # 优先 markitdown：保留标题/表格/列表结构；不可用或失败再回退 PyPDF2
    if file2md.is_available():
        try:
            md_text = file2md.file_to_markdown(path)
            if md_text and md_text.strip():
                return os.path.basename(path), md_text[:6000], None
        except Exception as e:  # noqa: BLE001
            common.logger.warning("markitdown 转 PDF 失败，回退 PyPDF2: %s", e)
    from PyPDF2 import PdfReader

    reader = PdfReader(path)
    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    return os.path.basename(path), text[:6000], len(reader.pages)


# 图片走 analyze_image（带视觉模型）；其余文档/数据/音频走 markitdown
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg"}
DOC_EXTS = {
    ".pdf", ".docx", ".pptx", ".xlsx", ".epub", ".html", ".htm",
    ".csv", ".json", ".xml", ".txt", ".md",
    ".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac",
}


def classify_file(path: str) -> str:
    """按扩展名把本地文件归类成素材类型：图片→image，其余→file(markitdown)。"""
    ext = os.path.splitext(path)[1].lower()
    if ext in IMAGE_EXTS:
        return "image"
    return "file"


def read_file_text(path: str):
    """用 markitdown 把文件转 Markdown。返回 (text, error)；失败则 (None, 原因)。"""
    if not file2md.is_available():
        return None, "markitdown 未安装（pip install 'markitdown[all]'）"
    try:
        md_text = file2md.file_to_markdown(path)
    except Exception as e:  # noqa: BLE001
        return None, f"markitdown 转换异常: {e}"
    if not md_text or not md_text.strip():
        return None, "markitdown 转换结果为空"
    return md_text[:6000], ""


def analyze_text(title: str, text: str, config: dict, extra: dict | None = None) -> dict:
    user = json.dumps({"title": title, "content": text}, ensure_ascii=False)
    try:
        res = common.llm_chat(MATERIAL_SYSTEM, user, config, expect_json=True)
    except Exception as e:  # noqa: BLE001
        common.logger.warning("分析失败: %s", e)
        res = {"title": title, "summary": (text[:200] + "..."), "tags": [], "key_points": []}
    if extra:
        res.update(extra)
    return res


def analyze_image(path: str, config: dict) -> dict:
    with open(path, "rb") as f:
        data = f.read()
    meta = {
        "filename": os.path.basename(path),
        "size": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
    }
    analysis: dict = {"title": meta["filename"], "summary": "", "tags": [], "key_points": []}
    vision_model = config.get("llm", {}).get("vision_model")
    if vision_model:
        b64 = base64.b64encode(data).decode()
        prompt = (
            "描述这张图片：主体、图中的文字、适合用作什么类型的新媒体素材。"
            '只输出 JSON：{"description":"","objects":[],"text_in_image":"","suitable_for":""}'
        )
        try:
            r = common.llm_chat(prompt, "请分析图片", config, expect_json=True, vision_b64=b64)
            analysis.update(r)
        except Exception as e:  # noqa: BLE001
            common.logger.warning("视觉分析失败: %s", e)
            analysis["summary"] = "（视觉模型不可用，仅记录元数据）"
    else:
        analysis["summary"] = "（未配置视觉模型，仅记录文件元数据）"
    analysis.update({"file_size": meta["size"], "sha256": meta["sha256"], "filename": meta["filename"]})
    return analysis


def to_markdown(a: dict) -> str:
    lines = [f"# {a.get('title', '素材')}", "", f"> 摘要：{a.get('summary', '')}", ""]
    if a.get("category"):
        lines.append(f"- 分类：{a['category']}")
    if a.get("tags"):
        lines.append(f"- 标签：{', '.join(a['tags'])}")
    if a.get("key_points"):
        lines.append("- 核心要点：")
        lines += [f"  - {p}" for p in a["key_points"]]
    if a.get("action_items"):
        lines.append("- 行动项：")
        lines += [f"  - {p}" for p in a["action_items"]]
    if a.get("reuse_suggestion"):
        lines.append(f"- 复用建议：{a['reuse_suggestion']}")
    return "\n".join(lines)


def write_docx(path: str, analysis: dict) -> None:
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(analysis.get("title", "素材"), level=1)
        doc.add_paragraph(analysis.get("summary", ""))
        if analysis.get("key_points"):
            doc.add_paragraph("核心要点：")
            for p in analysis["key_points"]:
                doc.add_paragraph(p, style="List Bullet")
        if analysis.get("action_items"):
            doc.add_paragraph("行动项：")
            for p in analysis["action_items"]:
                doc.add_paragraph(p, style="List Bullet")
        doc.save(path)
    except Exception as e:  # noqa: BLE001
        common.logger.debug("docx 导出跳过: %s", e)


def process_item(item: dict, config: dict, out_dir: str, dry_run: bool, yes: bool) -> dict | None:
    typ = item.get("type")
    raw = str(item.get("value", ""))
    key = common.dedup_key(raw) if typ != "text" else hashlib.sha1(item.get("value", "").encode("utf-8")).hexdigest()
    if key and key in common.load_seen(config):
        common.logger.info("已归档过，跳过: %s", raw[:60])
        return None
    if typ == "url":
        title, text = fetch_article(item["value"])
        analysis = analyze_text(title, text, config)
    elif typ == "pdf":
        title, text, pages = read_pdf(item["value"])
        analysis = analyze_text(title, text, config, extra={"page_count": pages})
    elif typ == "file":
        title = os.path.basename(item["value"])
        text, err = read_file_text(item["value"])
        if not text:
            # PDF 可回退 PyPDF2；其它格式无后备则跳过
            if item["value"].lower().endswith(".pdf"):
                t_title, text, pages = read_pdf(item["value"])
                analysis = analyze_text(t_title, text, config, extra={"page_count": pages, "convert_note": err})
            else:
                common.logger.warning("文件转换失败，跳过: %s (%s)", item["value"], err)
                return None
        else:
            analysis = analyze_text(title, text, config, extra={"source_file": title})
    elif typ == "image":
        analysis = analyze_image(item["value"], config)
    elif typ == "text":
        analysis = analyze_text(item.get("title", "文本素材"), item["value"], config)
    else:
        common.logger.warning("未知素材类型: %s", typ)
        return None

    # 摘要 / 要点 / 行动项经 renhua 去 AI 味（未启用或无 key 时原样保留）
    for fld in ("summary", "reuse_suggestion"):
        if analysis.get(fld):
            analysis[fld] = common.polish_text(analysis[fld], config)
    for fld in ("key_points", "action_items"):
        lst = analysis.get(fld) or []
        if lst:
            joined = common.polish_text("\n".join(lst), config)
            analysis[fld] = [k for k in joined.split("\n") if k.strip()]

    md = to_markdown(analysis)
    safe = "".join(c if c.isalnum() else "_" for c in (analysis.get("title") or "material"))[:40]
    md_path = os.path.join(out_dir, f"{safe}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)
    write_docx(os.path.join(out_dir, f"{safe}.docx"), analysis)

    doc_res = common.create_doc(analysis.get("title", "素材"), md, config, dry_run=dry_run)
    doc_url = doc_res.get("data", {}).get("url") if not dry_run else ""
    feishu = config.get("feishu", {})
    materials_table = feishu.get("materials_table_id")
    record = {
        "fields": {
            "标题": analysis.get("title", ""),
            "分类": analysis.get("category", "其他"),
            "标签": ", ".join(analysis.get("tags") or []),
            "文档链接": doc_url,
            "本地文件": md_path,
            "类型": typ,
        }
    }
    if materials_table and not dry_run:
        common.write_bitable_records(feishu.get("app_token"), materials_table, [record], config, dry_run=dry_run, yes=yes)
    if not dry_run and key:
        common.mark_seen(config, key)
    return {"title": analysis.get("title"), "md": md_path, "doc_url": doc_url}


def main() -> None:
    p = argparse.ArgumentParser(description="场景C：多模态素材管理")
    p.add_argument("--url")
    p.add_argument("--image")
    p.add_argument("--pdf")
    p.add_argument("--file", help="本地文件入库：自动按扩展名识别 docx/pptx/xlsx/epub/html/音频等 → markitdown 转 Markdown → 分析")
    p.add_argument("--text")
    p.add_argument("--queue", help="批量处理：JSON 文件，元素为 {type,value}")
    p.add_argument("--config")
    p.add_argument("--dry-run", action="store_true")
    p.add_argument("--yes", action="store_true")
    p.add_argument("--no-notify", action="store_true")
    p.add_argument("-v", "--verbose", action="store_true")
    args = p.parse_args()

    common.setup_logging(args.verbose)
    config = common.load_config(args.config)
    out_dir = common.ensure_output_dir(config)

    items: list[dict] = []
    if args.url:
        items.append({"type": "url", "value": args.url})
    if args.image:
        items.append({"type": "image", "value": args.image})
    if args.pdf:
        items.append({"type": "pdf", "value": args.pdf})
    if args.file:
        ftype = classify_file(args.file)
        items.append({"type": ftype, "value": args.file})
    if args.text:
        items.append({"type": "text", "value": args.text})
    if args.queue:
        with open(args.queue, "r", encoding="utf-8") as f:
            items += json.load(f)
    if not items:
        sys.exit("请指定 --url / --image / --pdf / --file / --text / --queue 之一。")

    results: list[dict] = []
    for it in items:
        r = process_item(it, config, out_dir, args.dry_run, args.yes)
        if r:
            results.append(r)
            common.logger.info("已处理: %s -> %s", r["title"], r["md"])

    feishu = config.get("feishu", {})
    if not args.no_notify and not args.dry_run and feishu.get("chat_id"):
        common.send_im(feishu["chat_id"], f"素材归档完成：{len(results)} 个", config, as_bot=True)

    print(json.dumps({"ok": True, "count": len(results), "results": results}, ensure_ascii=False))


if __name__ == "__main__":
    main()
